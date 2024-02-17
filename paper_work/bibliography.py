from docx import Document
import os
from pdf2docx import Converter
from django.core.exceptions import ObjectDoesNotExist
from bookshelf.source_citation import get_source_reference
from file_handling.models import PaperFile
from .models import Bibliography, Paper


def create_bibliography(paper: Paper):
    """Create new Bibliography obj based on new Paper obj"""
    new_bibliography = Bibliography(paper=paper)
    new_bibliography.save()


def update_bibliography(paper: Paper):
    """Update existing Bibliography obj based on chosen sources"""
 
    # Get apa and mla references
    bibliography = get_bibliography(paper)
    references = [get_source_reference(source) for source in paper.sources.all()]
    apa_sources: str = ""
    mla_sources: str = ""
    for number in range(len(references)):
        apa_sources += f"{number + 1}. {references[number].endnote_apa}\n\n"
        mla_sources += f"{number + 1}. {references[number].endnote_mla}\n\n"
    
    # Update obj
    bibliography.apa = apa_sources
    bibliography.mla = mla_sources
    return bibliography.save(update_fields=("apa", "mla",))


def clear_bibliography(paper: Paper):
    """Clear apa and mla fields of Bibliography obj"""
    bibliography = get_bibliography(paper)
    bibliography.apa, bibliography.mla = "", ""
    return bibliography.save(update_fields=("apa", "mla",))


def append_bibliography_to_file(file: PaperFile, bibliography: str):
    """Add bibliography to the end of a last uploaded file"""

    # If document extension is .pdf - convert it to .docx first
    if file.file_name().lower().endswith(".pdf"):

        # Get new file name and location
        future_file_name = f"{str(file.file_name().split('.')[0:-1][0])}.docx"
        output_path = os.path.join(file.get_directory_path(), future_file_name)

        # Convert .pdf to .docx
        converter = Converter(file.get_path_to_file())
        converter.convert(output_path)
        converter.close()
        path_to_file = output_path

        # Delete old .pdf file
        os.remove(file.get_path_to_file())

        # Update PaperFile obj (set file property to new .docx file)
        file.file = "/".join(output_path.split("/")[1::]) # delete media_root from file path
        file.save(update_fields=("file",))

    # If document extension is .docx - just get its location
    elif file.file_name().lower().endswith(".docx"):
        path_to_file = file.get_path_to_file()
    # Error case
    else:
        return None

    # Append bibliography
    file = Document(path_to_file)
    file.add_page_break()
    file.add_paragraph("Bibliography\n\n\n")
    file.add_paragraph(bibliography)
    return file.save(path_to_file)


def get_right_bibliography(paper: Paper) -> str | None:
    """Check if bibliography obj for existing paper obj exists"""

    bibliography = get_bibliography(paper)
    if not bibliography:
        return None
    if paper.citation_style == "APA":
        return bibliography.apa
    elif paper.citation_style == "MLA":
        return bibliography.mla
    else:
        return None
  

def get_bibliography(paper: Paper) -> Bibliography | None:
    """Check if bibliography obj for existing paper obj exists"""
    try:
        return Bibliography.objects.get(paper=paper)
    except ObjectDoesNotExist:
        return None
    
